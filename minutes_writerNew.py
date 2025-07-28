# minutes_writer.py — “出力形式” ブロックを除去した完全版
import re
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def set_meiryo_font(run):
    run.font.name = "Meiryo"
    run.font.element.rPr.rFonts.set(qn("w:eastAsia"), "Meiryo")

def write_minutes_section(word_file_path: str,
                          output_file_path: str,
                          replaced_transcription: str):
    """
    議事録本文を書き込む。

    1. '■ 議事' 見出し直後に文字列を挿入
    2. 議事冒頭のテンプレート指示

          【出力形式】
          [Speaker X] 発話内容
          [Speaker X] 発話内容
          ---

       を正規表現でまるごと削除
    3. 末尾に『以上』を右寄せで追加
    4. 全文フォントを Meiryo に統一
    """
    # -----------------------------------------------------------
    # (0) 余計なテンプレート指示を削除
    # -----------------------------------------------------------
    cleaned_text = re.sub(
        r"【出力形式】[\s\S]*?---\s*",
        "",                       # ← 該当ブロックを空文字へ
        replaced_transcription,
        flags=re.MULTILINE
    ).strip()

    doc = Document(word_file_path)

    # -----------------------------------------------------------
    # (1) 「■ 議事」段落を探す
    # -----------------------------------------------------------
    marker_idx = next(
        (i for i, p in enumerate(doc.paragraphs) if "■ 議事" in p.text),
        None
    )
    if marker_idx is None:
        print("[ERROR] 『■ 議事』の見出しが見つかりません。")
        return

    # (1') 見出し直後の空白行を削除
    i = marker_idx + 1
    while i < len(doc.paragraphs) and not doc.paragraphs[i].text.strip():
        doc.paragraphs[i].clear()
        i += 1

    # -----------------------------------------------------------
    # (2) 議事録本文を段落として挿入
    # -----------------------------------------------------------
    body_paras = []
    for line in filter(None, cleaned_text.splitlines()):
        p = doc.add_paragraph(line)
        body_paras.append(p)

    # 「■ 議事」の直後へスライス代入
    doc.paragraphs[marker_idx + 1 : marker_idx + 1] = body_paras

    # -----------------------------------------------------------
    # (3) 末尾に「以上」
    # -----------------------------------------------------------
    end_para = doc.add_paragraph("以上")
    end_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # -----------------------------------------------------------
    # (4) フォント統一
    # -----------------------------------------------------------
    for para in doc.paragraphs:
        for run in para.runs:
            set_meiryo_font(run)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        set_meiryo_font(run)

    # -----------------------------------------------------------
    # (5) 保存
    # -----------------------------------------------------------
    doc.save(output_file_path)
    print(f"[INFO] 議事録が保存されました → {output_file_path}")
