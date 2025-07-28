from docx import Document
import unicodedata
import re
import datetime

# ── ヘルパ ───────────────────────────────────────────────
def normalize_text(text: str) -> str:
    return unicodedata.normalize("NFKC", text).strip().rstrip("：:")

def clean_repeated_labels(label: str, value: str) -> str:
    label_norm = normalize_text(label)
    val = unicodedata.normalize("NFKC", str(value)).strip()
    # (a) 「- label:」「label:」等
    val = re.sub(rf"^[-\s]*{re.escape(label_norm)}[\s:：\-～ー]*", "", val).strip()
    # (b) 「1. - label:」
    return re.sub(rf"\d+\s*[\.．]\s*[-]?\s*{re.escape(label_norm)}\s*[:：]?", "", val).strip()

# ── メイン ───────────────────────────────────────────────
def table_writer(word_file_path: str, output_file_path: str, extracted_info: dict):
    doc = Document(word_file_path)

    # (1) テンプレ側ラベル → 正規化テーブル
    label_map = {}
    for tbl in doc.tables:
        for row in tbl.rows:
            if len(row.cells) < 2:
                continue
            raw = row.cells[0].text.strip()
            label_map[normalize_text(raw)] = raw
    print("[LOG] 正規化ラベル:", label_map)

    # (2) 抽出結果キーの正規化
    info_norm = {normalize_text(k): v for k, v in extracted_info.items()}

    # (3) 転記ループ
    for tbl in doc.tables:
        for row in tbl.rows:
            if len(row.cells) < 2:
                continue

            raw_label = row.cells[0].text.strip()
            norm_label = normalize_text(raw_label)

            if norm_label in info_norm:
                raw_value = info_norm[norm_label]

                # ---- 出席者：カンマ区切り化＆数字除去 --------------------
                if norm_label == normalize_text("出席者"):
                    parts = re.split(r"[、,;\n\r]+", str(raw_value))
                    cleaned_parts = []
                    for part in parts:
                        p = part.strip()
                        # 「1. 加藤」「１．加藤」「・1. 加藤」などを除去
                        p = re.sub(r"^[・●■]?", "", p)               # 先頭の黒点
                        p = re.sub(r"^\s*\d+\s*[\.．]\s*", "", p)       # 先頭番号
                        if p:
                            cleaned_parts.append(p)
                    raw_value = ", ".join(cleaned_parts)

                # ---- 次回日時：今年の西暦付与 ----------------------------
                if norm_label == normalize_text("次回会議予定日時"):
                    txt = str(raw_value)
                    year = datetime.datetime.now().year
                    raw_value = re.sub(
                        r"(\d{1,2})月\s*(\d{1,2})日",
                        lambda m: f"{year}年{int(m.group(1))}月{int(m.group(2))}日",
                        txt
                    )

                # ---- ラベル重複除去 ------------------------------------
                cleaned = clean_repeated_labels(raw_label, raw_value)

                # ---- 黒点付与（番号付き行は除外） ------------------------
                def add_bullet(line: str) -> str:
                    if not line.strip():
                        return line
                    # 既に「・」「●」等が付いている、または「1.」「1．」等で始まる行なら何もしない
                    if re.match(r"[・●■]", line) or re.match(r"\d+\s*[\.．]", line):
                        return line
                    return f"・{line}"

                bulleted = "\n".join(add_bullet(ln) for ln in cleaned.splitlines())

                row.cells[1].text = bulleted
                print(f"[LOG]  {raw_label} → {bulleted}")
            else:
                row.cells[1].text = ""
                print(f"[WARN] {raw_label}: データなし")

    # (4) 保存
    doc.save(output_file_path)
    print(f"[INFO] 保存完了 → {output_file_path}")
